"""
Document Scanner — Fast file inspection, zero LLM calls.

Inspects uploaded files and returns metadata needed for extraction:
  - file type, page count, is_scanned, has_toc

This is intentionally NOT an agent. It's pure deterministic Python.
No LLM calls are made here — LLM calls belong in DocumentAgent (extraction)
or the TutorAgent (teaching).

Usage:
    from agents.document_scanner_agent import scan, DocumentScannerAgent

    # Preferred: plain function
    entry = scan("physics.pdf")
    # returns DocumentScanEntry(file_name, file_path, file_type, total_pages,
    #                           is_scanned, has_toc, toc_entry_count)

    # Backward compat: class wrapper (same result)
    scanner = DocumentScannerAgent()
    report = scanner.scan(["physics.pdf", "lecture.pptx"])
"""

import logging
from pathlib import Path
from typing import List, Optional

from models.data_models import DocumentScanEntry, DocumentScanReport

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Core scan function — pure Python, zero LLM
# ---------------------------------------------------------------------------

def scan(file_path: str) -> Optional[DocumentScanEntry]:
    """
    Inspect a single file and return its metadata.

    Returns DocumentScanEntry or None if the file doesn't exist or is unsupported.
    Makes zero LLM calls.
    """
    path = Path(file_path)
    if not path.exists():
        logger.warning(f"File not found: {file_path}")
        return None

    ext = path.suffix.lower()
    size_kb = round(path.stat().st_size / 1024, 1)

    entry = DocumentScanEntry(
        file_name=path.name,
        file_path=str(path),
        file_type=_ext_to_type(ext),
        file_size_kb=size_kb,
        total_pages=0,
    )

    if ext == ".pdf":
        _scan_pdf(entry)
    elif ext == ".pptx":
        _scan_pptx(entry)
    elif ext == ".docx":
        _scan_docx(entry)
    elif ext in (".png", ".jpg", ".jpeg", ".webp"):
        _scan_image(entry)
    elif ext in (".txt", ".md"):
        _scan_text(entry)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return None

    logger.info(
        f"Scanned {entry.file_name}: type={entry.file_type}, "
        f"pages={entry.total_pages}, scanned={entry.is_scanned}, toc={entry.has_toc}"
    )
    return entry


# ---------------------------------------------------------------------------
# Per-format inspectors (no LLM)
# ---------------------------------------------------------------------------

def _scan_pdf(entry: DocumentScanEntry) -> None:
    """Inspect PDF using PyMuPDF — zero LLM calls."""
    from content.document_processor import classify_pdf

    result = classify_pdf(entry.file_path)
    entry.total_pages = result["total_pages"]
    entry.has_toc = result["has_toc"]
    entry.toc_entry_count = result["toc_count"]
    entry.is_scanned = result["is_scanned"]


def _scan_pptx(entry: DocumentScanEntry) -> None:
    """Inspect PowerPoint — zero LLM calls."""
    try:
        from pptx import Presentation
        prs = Presentation(entry.file_path)
        entry.total_pages = len(prs.slides)
        # PPTX slides always have a navigable structure (slide titles)
        entry.has_toc = entry.total_pages > 0
        entry.toc_entry_count = entry.total_pages
    except Exception as e:
        logger.debug(f"PPTX scan failed: {e}")


def _scan_docx(entry: DocumentScanEntry) -> None:
    """Inspect Word document — zero LLM calls."""
    try:
        from docx import Document
        doc = Document(entry.file_path)
        heading_count = sum(
            1 for p in doc.paragraphs
            if p.style and p.style.name.startswith("Heading")
        )
        entry.has_toc = heading_count > 3
        entry.toc_entry_count = heading_count
        entry.total_pages = max(1, len(doc.paragraphs) // 25)
    except Exception as e:
        logger.debug(f"DOCX scan failed: {e}")


def _scan_image(entry: DocumentScanEntry) -> None:
    """Mark image — always treated as scanned, 1 page."""
    entry.total_pages = 1
    entry.is_scanned = True


def _scan_text(entry: DocumentScanEntry) -> None:
    """Inspect text/markdown file."""
    try:
        size = Path(entry.file_path).stat().st_size
        entry.total_pages = max(1, size // 3000)
    except Exception:
        entry.total_pages = 1


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _ext_to_type(ext: str) -> str:
    mapping = {
        ".pdf": "pdf",
        ".pptx": "pptx",
        ".docx": "docx",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".webp": "image",
        ".txt": "text",
        ".md": "text",
    }
    return mapping.get(ext, "unknown")


# ---------------------------------------------------------------------------
# Backward-compatible class wrapper
# ---------------------------------------------------------------------------

class DocumentScannerAgent:
    """
    Thin wrapper around scan() for backward compatibility.

    The constructor still accepts llm_client for API compat, but it is ignored —
    scanning is now purely deterministic (zero LLM calls).
    """

    def __init__(self, llm_client=None):
        # llm_client kept in signature for backward compat; not used
        if llm_client is not None:
            logger.debug("DocumentScannerAgent: llm_client ignored (scanning is LLM-free)")

    def scan(self, file_paths: List[str]) -> DocumentScanReport:
        """
        Scan a list of files and return a DocumentScanReport.

        Zero LLM calls. Fast — pure Python file inspection only.
        """
        entries = []
        for path in file_paths:
            entry = scan(path)
            if entry:
                entries.append(entry)

        # Build a simple summary
        total_pages = sum(e.total_pages for e in entries)
        summary = f"{len(entries)} document(s) uploaded ({total_pages} total pages)"

        return DocumentScanReport(
            entries=entries,
            collection_summary=summary,
        )

    def scan_to_markdown(self, file_paths: List[str]) -> str:
        """Backward-compatible method: scan and return markdown table."""
        report = self.scan(file_paths)
        lines = [
            "| # | File | Type | Pages | Scanned | Has TOC |",
            "|---|------|------|-------|---------|---------|",
        ]
        for i, e in enumerate(report.entries, 1):
            lines.append(
                f"| {i} | {e.file_name} | {e.file_type} | {e.total_pages} "
                f"| {e.is_scanned} | {e.has_toc} |"
            )
        lines.append(f"\n{report.collection_summary}")
        return "\n".join(lines)
