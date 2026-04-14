"""
DOCX Processor — Microsoft Word documents.

Extracts text using heading hierarchy to build TOC.
Each Heading 1 = chapter, Heading 2 = section, etc.
"""

import logging
from pathlib import Path
from typing import List, Tuple

from models.data_models import ProcessedDocument, TOCEntry, Section

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Processes .docx files into ProcessedDocument format."""

    def process(self, file_path: str) -> ProcessedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")

        doc = Document(file_path)
        label = path.stem.replace("_", " ").replace("-", " ").title()

        # Group paragraphs into sections by heading level
        toc_entries, sections = self._build_sections(doc, file_path)

        return ProcessedDocument(
            file_path=file_path,
            file_type="docx",
            label=label,
            toc=toc_entries,
            sections=sections,
        )

    def _build_sections(self, doc, source_file: str) -> Tuple[List[TOCEntry], List[Section]]:
        toc_entries = []
        sections = []

        current_heading = None
        current_content_lines = []
        section_counter = 0
        pseudo_page = 1  # DOCX has no pages; use para count as proxy

        def flush_section():
            nonlocal pseudo_page
            if current_heading is None:
                return
            title, level, section_id = current_heading
            content = "\n".join(current_content_lines).strip()
            toc_entries.append(TOCEntry(
                title=title,
                level=level,
                page_or_slide=pseudo_page,
                source_file=source_file,
                section_id=section_id,
            ))
            sections.append(Section(
                section_id=section_id,
                title=title,
                content=content,
                page_start=pseudo_page,
                page_end=pseudo_page,
                source_file=source_file,
            ))

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if "Heading 1" in style_name:
                flush_section()
                section_counter += 1
                current_heading = (text, 1, f"h1_{section_counter}")
                current_content_lines = []
                pseudo_page += 1

            elif "Heading 2" in style_name:
                flush_section()
                section_counter += 1
                current_heading = (text, 2, f"h2_{section_counter}")
                current_content_lines = []
                pseudo_page += 1

            elif "Heading 3" in style_name:
                # Append to current section rather than creating new one
                current_content_lines.append(f"\n### {text}")

            else:
                current_content_lines.append(text)

        flush_section()

        # If no headings found, treat whole doc as one section
        if not sections:
            all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            toc_entries.append(TOCEntry(
                title="Document", level=1, page_or_slide=1,
                source_file=source_file, section_id="doc_full"
            ))
            sections.append(Section(
                section_id="doc_full", title="Document",
                content=all_text, page_start=1, page_end=1,
                source_file=source_file,
            ))

        return toc_entries, sections
